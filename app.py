import streamlit as st
from typing_extensions import TypedDict
from langgraph.graph import StateGraph, START, END
from langchain_core.tools import tool
from langchain_groq import ChatGroq
from langchain_tavily import TavilySearch
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

load_dotenv()

# TAVILY API & GROQ API setup
os.environ["TAVILY_API_KEY"]=os.getenv("TAVILY_API_KEY")
os.environ["GROQ_API_KEY"]=os.getenv("GROQ_API_KEY")
search_tool=TavilySearch(max_results=3)

# State schema
class State(TypedDict):
    topic: str
    research: str
    blog_draft: str
    needs_improvement: bool

@tool
def search_web(query: str) -> str:
    """Search the web for a query."""
    results=search_tool.invoke(query)
    return str(results)

def research_content(state: State):
    search_results=search_web.invoke(state["topic"])
    prompt=f"""Using the following search results, write a research summary for a blog post about {state['topic']}:

{search_results}

Provide:
1. A concise summary of key findings, recent developments, and actionable insights
2. List of sources with their titles and URLs for reference

Format your response as:
[Your summary here]

## Sources
[List the sources with titles and URLs]"""
    response=st.session_state["llm"].invoke(prompt)
    st.session_state["research_output"]=response.content
    return {"research": response.content}

def generate_blog_draft(state: State):
    draft_prompt=f"""Write an engaging blog post about {state['topic']} using this research:

{state['research']}

Make it informative, well-structured, and engaging for readers."""
    draft=st.session_state["llm"].invoke(draft_prompt)

    quality_prompt=f"""Evaluate this blog post strictly and respond with only "PASS" or "FAIL":

{draft.content}

STRICT CRITERIA (ALL must be met for PASS):
1. Length: At least 300 words
2. Structure: Has clear introduction, body, and conclusion
3. Engagement: Uses engaging headlines, examples, or stories
4. Value: Provides specific, actionable insights (not just generic advice)
5. Depth: Goes beyond surface-level information
6. Readability: Uses varied sentence structure and clear language

Count the word length and check each criterion carefully. If ANY criterion fails, respond "FAIL".
If the post is too short, generic, or lacks specific insights, respond "FAIL".

Your response:"""
    quality_check=st.session_state["llm"].invoke(quality_prompt)
    needs_improvement="FAIL" in quality_check.content.upper()
    st.session_state["draft_output"]=draft.content
    st.session_state["quality_check"]=quality_check.content
    return {
        "blog_draft": draft.content,
        "needs_improvement": needs_improvement
    }

def improve_blog_draft(state: State):
    improve_prompt=f"""Improve this blog post by:
1. Making it more engaging and readable
2. Adding more specific examples or insights
3. Improving structure and flow
4. Ensuring it provides clear value to readers

Original post:
{state['blog_draft']}

Provide the improved version:"""
    improved=st.session_state["llm"].invoke(improve_prompt)
    st.session_state["improved_output"]=improved.content
    return {"blog_draft": improved.content, "needs_improvement": False}

def route_after_draft(state: State):
    return "improve" if state["needs_improvement"] else END

# Streamlit UI
st.set_page_config(page_title="Blog Generator Squad", layout="centered")
st.title("Agentic AI Blog Generator")

groq_model=st.text_input("Enter Groq model (e.g. llama3-8b-8192 or mixtral-8x7b-32768)")

topic=st.text_input("Enter a blog topic:", placeholder="e.g. Agentic AI vs AI Agents")

if st.button("Generate Blog"):
    if not all([groq_model.strip(), topic.strip()]):
        st.warning("Please fill in all required fields.")
    else:
        from langchain_groq import ChatGroq
        st.session_state.clear()
        st.session_state["llm"]=ChatGroq(model=groq_model)

        # LangGraph pipeline setup
        workflow=StateGraph(State)
        workflow.add_node("researcher", research_content)
        workflow.add_node("writer", generate_blog_draft)
        workflow.add_node("editor", improve_blog_draft)
        workflow.add_edge(START, "researcher")
        workflow.add_edge("researcher", "writer")
        workflow.add_conditional_edges("writer", route_after_draft, {"improve": "editor", END: END})
        workflow.add_edge("editor", END)
        graph=workflow.compile()
        st.session_state.clear()
        st.session_state["llm"]=ChatGroq(model=groq_model)
        st.info("🧪 Researching and drafting...")
        result=graph.invoke({"topic": topic})
        st.subheader("🔍 Research Summary")
        st.markdown(st.session_state.get("research_output", "_No research generated._"))
        st.subheader("✍️ Initial Draft")
        st.markdown(st.session_state.get("draft_output", "_No draft generated._"))
        st.subheader("📊 Quality Check")
        st.markdown(f"**Result:** {st.session_state.get('quality_check', '_No quality check result._')}")
        if "improved_output" in st.session_state:
            st.subheader("🛠️ Edited Version")
            st.markdown(st.session_state["improved_output"])
        st.success("✅ Blog generation complete!")
        st.balloons()
        final_blog=st.session_state.get("improved_output") or st.session_state.get("draft_output")
        # Create Word doc with only the final blog post
        doc=Document()
        doc.add_heading(f"Blog: {topic}", level=1)
        doc.add_paragraph(final_blog)
        # Save to memory buffer
        buffer=BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # Streamlit download button
        st.download_button(
            label="📥 Download Final Blog as .docx",
            data=buffer,
            file_name=f"{topic.replace(' ','_')}_blog.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
